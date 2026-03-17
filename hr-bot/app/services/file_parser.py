"""
文件解析服务 - 支持PDF、Word、Excel文件解析
"""

import io
import logging
from typing import Optional
from pathlib import Path

logger = logging.getLogger(__name__)


class FileParser:
    """文件解析器 - 解析PDF、Word、Excel文件为文本"""

    @staticmethod
    def parse_pdf(file_content: bytes) -> str:
        """
        解析PDF文件

        Args:
            file_content: PDF文件二进制内容

        Returns:
            提取的文本内容
        """
        try:
            from pypdf import PdfReader

            pdf_file = io.BytesIO(file_content)
            reader = PdfReader(pdf_file)

            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"[FileParser] PDF解析失败: {e}")
            return ""

    @staticmethod
    def parse_docx(file_content: bytes) -> str:
        """
        解析Word文档 (docx)

        Args:
            file_content: Word文件二进制内容

        Returns:
            提取的文本内容
        """
        try:
            from docx import Document

            doc_file = io.BytesIO(file_content)
            doc = Document(doc_file)

            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # 也提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"[FileParser] DOCX解析失败: {e}")
            return ""

    @staticmethod
    def parse_doc(file_content: bytes) -> str:
        """
        解析旧版Word文档 (doc) - 使用antiword或olefile

        Args:
            file_content: Word文件二进制内容

        Returns:
            提取的文本内容
        """
        try:
            # 方法1: 尝试使用antiword命令行工具
            import subprocess
            import tempfile
            import os

            with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp:
                tmp.write(file_content)
                tmp_path = tmp.name

            try:
                result = subprocess.run(
                    ['antiword', tmp_path],
                    capture_output=True,
                    text=True,
                    timeout=10
                )
                if result.returncode == 0:
                    return result.stdout
            except (subprocess.SubprocessError, FileNotFoundError):
                pass

            # 方法2: 尝试使用olefile读取
            try:
                import olefile
                ole = olefile.OleFileIO(file_content)
                streams = ole.listdir()

                # 尝试读取WordDocument流
                if ['WordDocument'] in streams:
                    word_stream = ole.openstream('WordDocument')
                    data = word_stream.read()

                    # 简单的文本提取（去除控制字符）
                    text = ''
                    for i in range(0, len(data) - 1, 2):
                        try:
                            char = data[i:i+2].decode('utf-16-le', errors='ignore')
                            if char.isprintable() and char not in '\x00\x01\x02\x03\x04\x05\x06\x07\x08':
                                text += char
                        except:
                            continue
                    return text
            except ImportError:
                pass

            # 方法3: 尝试使用textract
            try:
                import textract
                with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp:
                    tmp.write(file_content)
                    tmp_path = tmp.name

                text = textract.process(tmp_path).decode('utf-8', errors='ignore')
                os.unlink(tmp_path)
                return text
            except (ImportError, Exception):
                pass

            # 清理临时文件
            try:
                os.unlink(tmp_path)
            except:
                pass

            logger.error("[FileParser] 无法解析DOC文件，请安装antiword或textract")
            return ""

        except Exception as e:
            logger.error(f"[FileParser] DOC解析失败: {e}")
            return ""

    @staticmethod
    def parse_excel(file_content: bytes) -> str:
        """
        解析Excel文件 (xlsx/xls)

        Args:
            file_content: Excel文件二进制内容

        Returns:
            提取的文本内容
        """
        try:
            import pandas as pd

            excel_file = io.BytesIO(file_content)

            # 读取所有sheet
            excel = pd.ExcelFile(excel_file)
            text_parts = []

            for sheet_name in excel.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)

                text_parts.append(f"=== Sheet: {sheet_name} ===")

                # 将DataFrame转换为文本
                for idx, row in df.iterrows():
                    row_text = []
                    for col in df.columns:
                        value = row[col]
                        if pd.notna(value):
                            row_text.append(f"{col}: {value}")
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"[FileParser] Excel解析失败: {e}")
            return ""

    @classmethod
    def parse_file(cls, file_content: bytes, filename: str) -> str:
        """
        根据文件类型自动解析

        Args:
            file_content: 文件二进制内容
            filename: 文件名（用于判断类型）

        Returns:
            提取的文本内容
        """
        suffix = Path(filename).suffix.lower()

        if suffix == '.pdf':
            return cls.parse_pdf(file_content)
        elif suffix == '.docx':
            return cls.parse_docx(file_content)
        elif suffix == '.doc':
            return cls.parse_doc(file_content)
        elif suffix in ['.xlsx', '.xls']:
            return cls.parse_excel(file_content)
        else:
            # 尝试作为文本读取
            try:
                return file_content.decode('utf-8')
            except UnicodeDecodeError:
                try:
                    return file_content.decode('gbk')
                except UnicodeDecodeError:
                    logger.error(f"[FileParser] 无法解析文件类型: {suffix}")
                    return ""


# 全局文件解析器实例
_file_parser: Optional[FileParser] = None


def get_file_parser() -> FileParser:
    """获取文件解析器单例"""
    global _file_parser
    if _file_parser is None:
        _file_parser = FileParser()
    return _file_parser
